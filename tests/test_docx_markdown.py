"""DOCX and Markdown ingestion.

DOCX is the format the target audience actually has - business documents in
Russian and Kazakh - and until now people had to print them to PDF first.
Markdown costs nothing: it goes through the text loader, since its syntax
carries meaning (a heading says what a section is about) rather than noise.

Fixtures are built in the test with python-docx, so the suite stays offline and
carries no binary files.
"""
import pytest
from docx import Document as Docx

from app.rag.document_loader import (
    SUPPORTED_EXTENSIONS,
    TEXT_EXTENSIONS,
    DocumentLoader,
)


def _docx(tmp_path, name="doc.docx", paragraphs=("Some prose.",), table=None):
    document = Docx()
    for text in paragraphs:
        document.add_paragraph(text)
    if table:
        created = document.add_table(rows=len(table), cols=len(table[0]))
        for row_index, row in enumerate(table):
            for column_index, value in enumerate(row):
                created.cell(row_index, column_index).text = value
    path = tmp_path / name
    document.save(path)
    return path


def _bytes_of(path):
    return path.read_bytes()


# =========================
# The format list is shared
# =========================

def test_the_extension_list_is_the_single_source():
    """The API gate and both clients read this, so a new format is added once."""
    assert ".docx" in SUPPORTED_EXTENSIONS
    assert ".md" in SUPPORTED_EXTENSIONS
    assert ".pdf" in SUPPORTED_EXTENSIONS
    assert all(extension.startswith(".") for extension in SUPPORTED_EXTENSIONS)


def test_text_extensions_are_a_subset():
    assert set(TEXT_EXTENSIONS) <= set(SUPPORTED_EXTENSIONS)


def test_nothing_hardcodes_a_format_list_any_more():
    """Four copies of "txt, pdf" is how a format gets half-added."""
    from pathlib import Path

    root = Path(__file__).resolve().parents[1]
    for relative in ("app/main.py", "frontend/streamlit_app.py", "clients/telegram_bot.py"):
        source = (root / relative).read_text(encoding="utf-8")
        assert "SUPPORTED_EXTENSIONS" in source, f"{relative} does not use the shared list"


# =========================
# DOCX
# =========================

def test_a_docx_is_loaded(tmp_path):
    path = _docx(tmp_path, paragraphs=("First paragraph.", "Second paragraph."))

    documents = DocumentLoader().load_document(str(path))

    assert len(documents) == 1
    assert "First paragraph." in documents[0].page_content
    assert "Second paragraph." in documents[0].page_content


def test_docx_metadata_names_the_format(tmp_path):
    path = _docx(tmp_path)

    metadata = DocumentLoader().load_document(str(path))[0].metadata

    assert metadata["type"] == "docx"
    assert metadata["source"] == "doc.docx"
    assert metadata["char_count"] > 0


def test_table_text_is_extracted(tmp_path):
    """The answer is usually in the table.

    python-docx does not include table text in document.paragraphs, so reading
    only paragraphs would index the prose around the answer and not the answer.
    """
    path = _docx(
        tmp_path,
        paragraphs=("Leave policy follows.",),
        table=[["Role", "Days"], ["Engineer", "28"], ["Director", "35"]],
    )

    content = DocumentLoader().load_document(str(path))[0].page_content

    assert "Engineer" in content and "28" in content
    assert "Director" in content and "35" in content


def test_table_rows_stay_together(tmp_path):
    """A row split across lines loses which number belongs to which role."""
    path = _docx(tmp_path, table=[["Engineer", "28"], ["Director", "35"]])

    content = DocumentLoader().load_document(str(path))[0].page_content

    assert "Engineer | 28" in content
    assert "Director | 35" in content


def test_a_merged_cell_is_not_repeated(tmp_path):
    """python-docx returns a merged cell once per column it spans, so a header
    spanning the table would arrive as "Leave | Leave | Leave"."""
    document = Docx()
    table = document.add_table(rows=2, cols=3)
    table.cell(0, 0).merge(table.cell(0, 2)).text = "Leave policy"
    for column, value in enumerate(("Role", "Days", "Notes")):
        table.cell(1, column).text = value
    path = tmp_path / "merged.docx"
    document.save(path)

    content = DocumentLoader().load_document(str(path))[0].page_content

    assert "Leave policy" in content
    assert "Leave policy | Leave policy" not in content
    assert "Role | Days | Notes" in content, content


def test_the_table_count_is_recorded(tmp_path):
    path = _docx(tmp_path, table=[["a", "b"]])

    assert DocumentLoader().load_document(str(path))[0].metadata["tables"] == 1


def test_cyrillic_content_survives(tmp_path):
    """The audience this format was added for."""
    path = _docx(
        tmp_path,
        paragraphs=("Ежегодный отпуск предоставляется по заявлению.",),
        table=[["Должность", "Дней"], ["Инженер", "28"]],
    )

    content = DocumentLoader().load_document(str(path))[0].page_content

    assert "Ежегодный отпуск" in content
    assert "Инженер | 28" in content


def test_an_empty_docx_is_rejected(tmp_path):
    path = _docx(tmp_path, paragraphs=())

    with pytest.raises(ValueError, match="no text"):
        DocumentLoader().load_document(str(path))


def test_a_whitespace_only_docx_is_rejected(tmp_path):
    path = _docx(tmp_path, paragraphs=("   ", "\t"))

    with pytest.raises(ValueError, match="no text"):
        DocumentLoader().load_document(str(path))


def test_a_corrupt_docx_raises(tmp_path):
    path = tmp_path / "broken.docx"
    path.write_bytes(b"this is definitely not a docx")

    with pytest.raises(Exception):
        DocumentLoader().load_document(str(path))


def test_a_missing_docx_raises(tmp_path):
    with pytest.raises(FileNotFoundError):
        DocumentLoader.load_docx(str(tmp_path / "nope.docx"))


# =========================
# Markdown
# =========================

@pytest.mark.parametrize("name", ["notes.md", "notes.markdown"])
def test_markdown_is_loaded(tmp_path, name):
    path = tmp_path / name
    path.write_text("# Title\n\n- first item\n- second item\n", encoding="utf-8")

    documents = DocumentLoader().load_document(str(path))

    assert "first item" in documents[0].page_content


def test_markdown_syntax_is_kept(tmp_path):
    """Headings and list markers say what a passage is; stripping them to bare
    prose would lose structure the model can use."""
    path = tmp_path / "notes.md"
    path.write_text("# Leave policy\n\n- 28 days\n", encoding="utf-8")

    content = DocumentLoader().load_document(str(path))[0].page_content

    assert "# Leave policy" in content
    assert "- 28 days" in content


def test_markdown_metadata_says_markdown_not_text(tmp_path):
    """The document listing shows this, so "txt" for a .md file would lie."""
    path = tmp_path / "notes.md"
    path.write_text("# Title\n", encoding="utf-8")

    assert DocumentLoader().load_document(str(path))[0].metadata["type"] == "md"


def test_a_plain_text_file_is_still_txt(tmp_path):
    path = tmp_path / "notes.txt"
    path.write_text("plain", encoding="utf-8")

    assert DocumentLoader().load_document(str(path))[0].metadata["type"] == "txt"


def test_markdown_in_a_legacy_encoding_is_decoded(tmp_path):
    """Markdown goes through the text loader, charset detection included.

    The sample is a few sentences long on purpose: charset detection is
    statistical, and on a one-line file cp1251 Cyrillic is genuinely
    indistinguishable from other byte-pair encodings.
    """
    text = (
        "# Отпуск\n\n"
        "Ежегодный оплачиваемый отпуск предоставляется по заявлению сотрудника.\n"
        "- Инженер: двадцать восемь календарных дней\n"
        "- Директор: тридцать пять календарных дней\n"
    )
    path = tmp_path / "ru.md"
    path.write_bytes(text.encode("cp1251"))

    content = DocumentLoader().load_document(str(path))[0].page_content

    assert "Отпуск" in content
    assert "Инженер" in content


# =========================
# The old .doc format
# =========================

def test_the_old_doc_format_says_what_to_do(tmp_path):
    """People will try it. "Unsupported format" does not help them."""
    path = tmp_path / "legacy.doc"
    path.write_bytes(b"\xd0\xcf\x11\xe0 old binary format")

    with pytest.raises(ValueError) as exc:
        DocumentLoader().load_document(str(path))

    message = str(exc.value)
    assert ".docx" in message, message
    assert "PDF" in message, message


def test_an_unknown_format_lists_what_is_supported(tmp_path):
    path = tmp_path / "sheet.xlsx"
    path.write_bytes(b"whatever")

    with pytest.raises(ValueError) as exc:
        DocumentLoader().load_document(str(path))

    for extension in SUPPORTED_EXTENSIONS:
        assert extension in str(exc.value)


# =========================
# Through the API
# =========================

def test_uploading_a_docx_indexes_it(api, tmp_path):
    path = _docx(
        tmp_path,
        paragraphs=("Annual leave is twenty eight days.",),
        table=[["Role", "Days"], ["Engineer", "28"]],
    )

    response = api.post(
        "/upload",
        params={"user_id": "u1"},
        files={"file": ("policy.docx", _bytes_of(path),
                        "application/vnd.openxmlformats-officedocument.wordprocessingml.document")},
    )

    assert response.status_code == 200, response.text
    assert response.json()["chunks"] >= 1


def test_a_docx_shows_up_in_the_document_list(api, tmp_path):
    path = _docx(tmp_path, paragraphs=("Content here.",))

    api.post(
        "/upload", params={"user_id": "u1"},
        files={"file": ("policy.docx", _bytes_of(path), "application/octet-stream")},
    )
    listing = api.get("/documents", params={"user_id": "u1"}).json()

    assert listing["documents"][0]["type"] == "docx"


def test_uploading_markdown_indexes_it(api):
    response = api.post(
        "/upload",
        params={"user_id": "u1"},
        files={"file": ("notes.md", b"# Notes\n\nRAG grounds answers.", "text/markdown")},
    )

    assert response.status_code == 200, response.text
    assert response.json()["chunks"] >= 1


def test_a_corrupt_docx_upload_is_a_400(api):
    response = api.post(
        "/upload",
        params={"user_id": "u1"},
        files={"file": ("broken.docx", b"not a docx at all", "application/octet-stream")},
    )

    assert response.status_code == 400
    assert [p for p in api.upload_dir.rglob("*") if p.is_file()] == []


def test_an_old_doc_upload_is_rejected_with_advice(api):
    """The gate rejects .doc before the loader sees it, so the message comes
    from the extension list - still better than silence."""
    response = api.post(
        "/upload",
        params={"user_id": "u1"},
        files={"file": ("legacy.doc", b"\xd0\xcf\x11\xe0 binary", "application/msword")},
    )

    assert response.status_code == 400
    assert "DOCX" in response.json()["detail"]


def test_the_rejection_message_lists_the_real_formats(api):
    response = api.post(
        "/upload",
        params={"user_id": "u1"},
        files={"file": ("sheet.xlsx", b"whatever", "application/octet-stream")},
    )

    detail = response.json()["detail"]
    assert "PDF" in detail and "DOCX" in detail and "MD" in detail
