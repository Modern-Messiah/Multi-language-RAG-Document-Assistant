"""
Document Loader for RAG Assistant
Supports PDF, DOCX, Markdown and plain text files with metadata extraction
"""

import logging
from pathlib import Path
from typing import List

from charset_normalizer import from_bytes
from docx import Document as DocxDocument
from langchain.schema import Document
from langchain_community.document_loaders import PyPDFLoader

logger = logging.getLogger(__name__)

# Read through load_txt, with charset detection for legacy encodings.
TEXT_EXTENSIONS = (".txt", ".md", ".markdown")

# The single source of truth for what may be uploaded. The API's extension gate
# and both clients' file pickers read this, so a new format is added in one
# place instead of four.
SUPPORTED_EXTENSIONS = (".pdf", ".docx", *TEXT_EXTENSIONS)


class DocumentLoader:
    """
    Load and parse documents with metadata extraction
    
    Supported formats: see SUPPORTED_EXTENSIONS above, which the API's upload
    gate and both clients read as well.
    - PDF (.pdf), one document per page
    - Word (.docx), paragraphs and table cells
    - Text (.txt) and Markdown (.md, .markdown), with charset detection

    Each document is loaded with metadata including:
    - source: filename
    - page: page number (PDF only)
    - total_pages: total pages in document (PDF only)
    - tables: number of tables read (DOCX only)
    - type: document type (pdf/docx/txt/md/markdown)
    """
    
    @staticmethod
    def load_pdf(file_path: str) -> List[Document]:
        """
        Load PDF document and extract text with page numbers
        
        Args:
            file_path: Path to PDF file
            
        Returns:
            List of Document objects, one per page
            
        Raises:
            FileNotFoundError: If file doesn't exist
            Exception: If PDF parsing fails
        """
        file_path = Path(file_path)
        
        if not file_path.exists():
            raise FileNotFoundError(f"PDF file not found: {file_path}")
        
        try:
            logger.info(f"📄 Loading PDF: {file_path.name}")
            
            loader = PyPDFLoader(str(file_path))
            documents = loader.load()
            
            if not documents:
                raise ValueError(f"PDF appears to be empty: {file_path.name}")
            
            # Enrich metadata
            for i, doc in enumerate(documents):
                doc.metadata.update({
                    "page": i + 1,
                    "source": file_path.name,
                    "total_pages": len(documents),
                    "type": "pdf",
                    "file_path": str(file_path)
                })
            
            logger.info(f"✅ Loaded {len(documents)} pages from {file_path.name}")
            return documents
            
        except Exception as e:
            logger.error(f"❌ Error loading PDF {file_path.name}: {str(e)}")
            raise
    
    @staticmethod
    def load_txt(file_path: str) -> List[Document]:
        """
        Load text document
        
        Args:
            file_path: Path to TXT file
            
        Returns:
            List containing single Document object
            
        Raises:
            FileNotFoundError: If file doesn't exist
            Exception: If text parsing fails
        """
        file_path = Path(file_path)
        
        if not file_path.exists():
            raise FileNotFoundError(f"TXT file not found: {file_path}")
        
        try:
            logger.info(f"📝 Loading TXT: {file_path.name}")

            raw = file_path.read_bytes()
            try:
                # utf-8-sig also handles plain utf-8 and strips a BOM if present
                text = raw.decode("utf-8-sig")
                encoding = "utf-8"
            except UnicodeDecodeError:
                # Legacy encodings (cp1251 for Russian/Kazakh, cp1252, koi8-r, ...)
                # cannot be told apart by decode-success, so use charset detection
                best = from_bytes(raw).best()
                if best is None:
                    raise ValueError(
                        f"Could not decode text file: {file_path.name}"
                    )
                text = str(best)
                encoding = best.encoding

            if not text.strip():
                raise ValueError(f"Text file appears to be empty: {file_path.name}")

            documents = [Document(
                page_content=text,
                metadata={
                    "source": file_path.name,
                    # The real extension, not "txt": Markdown also comes
                    # through here, and the document listing shows this.
                    "type": file_path.suffix.lower().lstrip(".") or "txt",
                    "file_path": str(file_path),
                    "char_count": len(text),
                    "encoding": encoding
                }
            )]
            
            logger.info(
                f"✅ Loaded {len(documents[0].page_content)} characters "
                f"from {file_path.name}"
            )
            return documents
            
        except Exception as e:
            logger.error(f"❌ Error loading TXT {file_path.name}: {str(e)}")
            raise
    
    @staticmethod
    def load_docx(file_path: str) -> List[Document]:
        """Load a .docx, including the text inside its tables.

        Tables matter more than they look: in the business documents this is
        aimed at, the rates, dates and headcounts a person asks about are
        usually in a table, and python-docx does not include table text in
        document.paragraphs. Reading only paragraphs would index the prose
        around the answer and not the answer.

        Headers, footers and footnotes are not read. Saying so is better than
        implying a completeness that is not there.
        """
        file_path = Path(file_path)

        if not file_path.exists():
            raise FileNotFoundError(f"DOCX file not found: {file_path}")

        try:
            logger.info(f"Loading DOCX: {file_path.name}")

            document = DocxDocument(str(file_path))

            blocks = [p.text.strip() for p in document.paragraphs if p.text.strip()]

            for table in document.tables:
                for row in table.rows:
                    cells = [cell.text.strip() for cell in row.cells]
                    # Deduplicate merged cells, which python-docx repeats once
                    # per column they span.
                    unique = []
                    for cell in cells:
                        if cell and (not unique or cell != unique[-1]):
                            unique.append(cell)
                    if unique:
                        blocks.append(" | ".join(unique))

            text = "\n\n".join(blocks)

            if not text.strip():
                raise ValueError(f"DOCX appears to contain no text: {file_path.name}")

            documents = [Document(
                page_content=text,
                metadata={
                    "source": file_path.name,
                    "type": "docx",
                    "file_path": str(file_path),
                    "char_count": len(text),
                    "paragraphs": len(document.paragraphs),
                    "tables": len(document.tables),
                }
            )]

            logger.info(
                f"Loaded {len(text)} characters from {file_path.name} "
                f"({len(document.tables)} table(s))"
            )
            return documents

        except Exception as e:
            logger.error(f"Error loading DOCX {file_path.name}: {str(e)}")
            raise

    @classmethod
    def load_document(cls, file_path: str) -> List[Document]:
        """
        Auto-detect format and load document
        
        Args:
            file_path: Path to document file
            
        Returns:
            List of Document objects
            
        Raises:
            FileNotFoundError: If file doesn't exist
            ValueError: If file format not supported
            Exception: If loading fails
        """
        file_path = Path(file_path)
        
        if not file_path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")
        
        extension = file_path.suffix.lower()

        if extension == ".pdf":
            return cls.load_pdf(str(file_path))
        if extension == ".docx":
            return cls.load_docx(str(file_path))
        if extension in TEXT_EXTENSIONS:
            # Markdown goes through the text loader unchanged. Its syntax
            # carries meaning a reader relies on - headings say what a section
            # is about, list items keep items apart - so stripping it to bare
            # prose would lose structure the model can use.
            return cls.load_txt(str(file_path))
        if extension == ".doc":
            # The pre-2007 binary format, which python-docx cannot read. People
            # will try it, so say what to do instead of "unsupported format".
            raise ValueError(
                "The old .doc format is not supported. Open it and save as "
                ".docx, or export it to PDF."
            )

        raise ValueError(
            f"Unsupported file format: {extension}\n"
            f"Supported formats: {', '.join(SUPPORTED_EXTENSIONS)}"
        )
    
    @staticmethod
    def get_document_info(documents: List[Document]) -> dict:
        """
        Get summary information about loaded documents
        
        Args:
            documents: List of Document objects
            
        Returns:
            Dictionary with document statistics
        """
        if not documents:
            return {"error": "No documents provided"}
        
        total_chars = sum(len(doc.page_content) for doc in documents)
        doc_type = documents[0].metadata.get("type", "unknown")
        
        info = {
            "num_documents": len(documents),
            "type": doc_type,
            "total_characters": total_chars,
            "source": documents[0].metadata.get("source", "unknown")
        }
        
        if doc_type == "pdf":
            info["total_pages"] = documents[0].metadata.get("total_pages", len(documents))
        
        return info


# Testing and example usage
if __name__ == "__main__":
    """
    Test the DocumentLoader with sample files
    """
    import tempfile
    
    print("\n" + "="*60)
    print("Testing DocumentLoader")
    print("="*60 + "\n")
    
    # Create a temporary test TXT file
    with tempfile.NamedTemporaryFile(mode='w', suffix='.txt', delete=False) as f:
        test_txt_path = f.name
        f.write("""
RAG (Retrieval-Augmented Generation) System Overview

RAG is a powerful technique that combines:
1. Information Retrieval: Finding relevant documents
2. Language Generation: Creating coherent responses

Key Components:
- Document Loader: Parses PDF/TXT files
- Text Splitter: Breaks text into chunks
- Embeddings: Converts text to vectors
- Vector Store: Stores and searches embeddings
- LLM Chain: Generates answers from context

Benefits:
- Reduces hallucinations
- Provides source attribution
- Works with custom knowledge bases
- More accurate than pure LLM generation

This system is built with LangChain, ChromaDB, and FastAPI.
        """.strip())
    
    try:
        # Test loading
        loader = DocumentLoader()
        
        print("1️⃣ Testing TXT loading...")
        print("-" * 60)
        
        docs = loader.load_document(test_txt_path)
        
        print("✅ Successfully loaded document")
        print(f"   Source: {docs[0].metadata['source']}")
        print(f"   Type: {docs[0].metadata['type']}")
        print(f"   Characters: {docs[0].metadata['char_count']}")
        
        # Get document info
        info = DocumentLoader.get_document_info(docs)
        print("\n📊 Document Info:")
        for key, value in info.items():
            print(f"   {key}: {value}")
        
        # Show content preview
        print("\n📄 Content Preview (first 200 chars):")
        print(f"   {docs[0].page_content[:200]}...")
        
        print("\n" + "="*60)
        print("✅ All tests passed!")
        print("="*60 + "\n")
        
        print("💡 Next steps:")
        print("   1. Create a PDF file to test PDF loading")
        print("   2. Try loading with: loader.load_document('your_file.pdf')")
        print("   3. Move on to Day 2: Text Chunking\n")
        
    except Exception as e:
        print(f"\n❌ Test failed: {str(e)}\n")
    
    finally:
        # Cleanup
        import os
        if os.path.exists(test_txt_path):
            os.unlink(test_txt_path)